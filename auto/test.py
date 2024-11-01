import pandas as pd
from docx import Document
import os

# 替换段落中的文本并保留字体样式
def replace_text_in_paragraph(paragraph, data):
    new_runs = []
    for run in paragraph.runs:
        run_text = run.text
        for key, value in data.items():
            if key in run_text:
                parts = run_text.split(key)
                for i in range(len(parts)):
                    new_runs.append(paragraph.add_run(parts[i]))
                    if i < len(parts) - 1:
                        new_run = paragraph.add_run(value)
                        # 设置新字体及样式
                        new_run.font.name = "標楷體"
                        new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
            else:
                new_runs.append(paragraph.add_run(run_text))

    # 清空段落内容
    paragraph.clear()
    # 将新的 run 添加回段落
    for new_run in new_runs:
        run_added = paragraph.add_run(new_run.text)
        run_added.font.name = new_run.font.name
        run_added.font.size = new_run.font.size
        run_added.font.bold = new_run.font.bold
        run_added.font.italic = new_run.font.italic

# 填充模板
def fill_template(template_path, output_path, data):
    if not os.path.exists(template_path):
        print(f"模板文件不存在: {template_path}")
        return
    
    try:
        doc = Document(template_path)
        print(f"正在填充模板: {template_path}")
        
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, data)

        for section in doc.sections:
            for header in section.header.paragraphs:
                replace_text_in_paragraph(header, data)
            for footer in section.footer.paragraphs:
                replace_text_in_paragraph(footer, data)

        doc.save(output_path)
        print(f"已生成文件: {output_path}")
    except Exception as e:
        print(f"填充模板时发生错误: {e}")

# 从 Excel 读取数据
def read_data_from_excel(file_path):
    try:
        df = pd.read_excel(file_path)
        print("读取到的数据：")
        print(df)
        data_list = []
        for _, row in df.iterrows():
            data = {
                "{會議名稱}": row["會議名稱"],
                "{會議字號}": row["會議字號"],
                "{會議主席}": row["會議主席"],
                "{會議記錄}": row["會議記錄"],
                "{開會日期}": row["開會日期"],
                "{開會地點}": row["開會地點"]
            }
            data_list.append(data)
        return data_list
    except Exception as e:
        print(f"读取 Excel 文件时发生错误: {e}")
        return []

# 主函数
def main():
    excel_file = "會議資料.xlsx"
    templates = {
        "成會單": "成會單模板.docx",
        "決議單": "決議單模板.docx",
        "提案單": "提案單模板.docx"
    }
    output_folder = "生成的會議文件"

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"创建输出文件夹: {output_folder}")

    data_list = read_data_from_excel(excel_file)
    
    for i, data in enumerate(data_list):
        for template_name, template_path in templates.items():
            output_path = os.path.join(output_folder, f"{data['{會議名稱}']}_{template_name}_{i+1}.docx")
            fill_template(template_path, output_path, data)

if __name__ == "__main__":
    main()
