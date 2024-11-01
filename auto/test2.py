from docx import Document

# 保留每个 run 的原始格式，仅在替换的时候设置字体
def replace_text_in_paragraph(paragraph, data):
    new_runs = []  # 保存新创建的 run

    for run in paragraph.runs:
        run_text = run.text
        for key, value in data.items():
            if key in run_text:
                parts = run_text.split(key)
                # 添加替换前的部分
                for i in range(len(parts)):
                    new_runs.append(paragraph.add_run(parts[i]))
                    # 如果不是最后一部分，则添加替换后的部分
                    if i < len(parts) - 1:
                        new_run = paragraph.add_run(value)
                        new_run.font.name = "標楷體"  # 设置新字体
                        new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
            else:
                new_runs.append(paragraph.add_run(run_text))
    
    # 清空段落
    paragraph.clear()
    # 将新创建的内容重新添加到段落
    for new_run in new_runs:
        paragraph.add_run(new_run.text).font = new_run.font

# 主函数，处理段落和表格
def fill_template_with_fonts(template_path, output_path, data):
    doc = Document(template_path)
    
    # 替换正文段落
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, data)

    # 替换表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)

    # 替换页首
    for section in doc.sections:
        for header in section.header.paragraphs:
            replace_text_in_paragraph(header, data)

    # 替换页尾
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            replace_text_in_paragraph(footer, data)

    # 保存结果文件
    doc.save(output_path)

# 测试用数据示例
data = {
    "{會議名稱}": "113年 九月份第一次經費審查會議",
    "{會議字號}": "SC-190009",
    "{會議主席}": "李承哲",
    "{會議記錄}": "陳翊嘉",
    "{開會日期}": "113年09月20日20:00",
    "{開會地點}": "中正館三樓公共空間"
}

fill_template_with_fonts("成會單模板.docx", "output.docx", data)
